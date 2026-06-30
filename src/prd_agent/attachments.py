from __future__ import annotations

from dataclasses import dataclass
from hashlib import sha256
from html.parser import HTMLParser
from pathlib import Path
import re
from typing import Any
from uuid import uuid4


TEXT_EXTENSIONS = {".md", ".txt"}
HTML_EXTENSIONS = {".html", ".htm"}
PDF_EXTENSIONS = {".pdf"}
DOCX_EXTENSIONS = {".docx"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp"}
ALLOWED_EXTENSIONS = (
    TEXT_EXTENSIONS
    | HTML_EXTENSIONS
    | PDF_EXTENSIONS
    | DOCX_EXTENSIONS
    | IMAGE_EXTENSIONS
)


@dataclass(frozen=True)
class IncomingAttachment:
    filename: str
    content_type: str | None
    content: bytes


class _TextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []
        self._skip_depth = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in {"script", "style"}:
            self._skip_depth += 1
        if tag in {"p", "br", "div", "li", "section", "article", "h1", "h2", "h3"}:
            self.parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in {"script", "style"} and self._skip_depth:
            self._skip_depth -= 1
        if tag in {"p", "div", "li", "section", "article"}:
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        if not self._skip_depth:
            self.parts.append(data)

    def text(self) -> str:
        return _compact_text(" ".join(self.parts))


def attachment_kind(filename: str) -> str:
    suffix = Path(filename).suffix.casefold()
    if suffix in TEXT_EXTENSIONS:
        return "text"
    if suffix in HTML_EXTENSIONS:
        return "html"
    if suffix in PDF_EXTENSIONS:
        return "pdf"
    if suffix in DOCX_EXTENSIONS:
        return "docx"
    if suffix in IMAGE_EXTENSIONS:
        return "image"
    raise ValueError(f"不支持的附件类型：{filename}")


def store_incoming_attachment(
    upload_dir: Path,
    project_id: str,
    attachment: IncomingAttachment,
) -> dict[str, Any]:
    raw_filename = Path(attachment.filename or "attachment").name.strip()
    if not raw_filename:
        raw_filename = "attachment"
    kind = attachment_kind(raw_filename)
    original_filename = _trim_filename(raw_filename, 255)
    safe_filename = _safe_filename(raw_filename)
    attachment_id = str(uuid4())
    stored_path = f"{project_id}/{attachment_id}-{safe_filename}"
    path = resolve_attachment_path(upload_dir, stored_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(attachment.content)
    return {
        "id": attachment_id,
        "original_filename": original_filename,
        "stored_path": stored_path,
        "content_type": attachment.content_type,
        "size_bytes": len(attachment.content),
        "sha256": sha256(attachment.content).hexdigest(),
        "kind": kind,
    }


def resolve_attachment_path(upload_dir: Path, stored_path: str) -> Path:
    root = upload_dir.resolve()
    path = (root / stored_path).resolve()
    if root != path and root not in path.parents:
        raise ValueError("附件路径非法")
    return path


def extract_attachment_text(
    attachment: Any,
    upload_dir: Path,
    llm: Any | None = None,
) -> str:
    path = resolve_attachment_path(upload_dir, str(attachment.stored_path))
    if not path.is_file():
        raise ValueError(f"附件文件不存在：{attachment.original_filename}")

    kind = str(attachment.kind)
    if kind == "text":
        text = path.read_text(encoding="utf-8")
    elif kind == "html":
        text = _extract_html_text(path)
    elif kind == "pdf":
        text = _extract_pdf_text(path)
    elif kind == "docx":
        text = _extract_docx_text(path)
    elif kind == "image":
        text = _analyze_image(path, str(attachment.original_filename), llm)
    else:
        raise ValueError(f"不支持的附件类型：{attachment.original_filename}")

    text = text.strip()
    if not text:
        raise ValueError(f"附件没有可用文本：{attachment.original_filename}")
    return text


def format_attachment_source(attachment: Any, extracted_text: str) -> str:
    return "\n".join(
        [
            f"附件名称：{attachment.original_filename}",
            f"附件类型：{attachment.kind}",
            "",
            extracted_text.strip(),
        ]
    )


def _safe_filename(filename: str) -> str:
    name = Path(filename or "attachment").name.strip()
    name = re.sub(r"[^A-Za-z0-9._ -]+", "_", name)
    name = name.strip(" .")
    if not name:
        name = "attachment"
    if Path(name).suffix.casefold() not in ALLOWED_EXTENSIONS:
        raise ValueError(f"不支持的附件类型：{filename}")
    return _trim_filename(name, 180)


def _trim_filename(filename: str, max_length: int) -> str:
    if len(filename) <= max_length:
        return filename
    path = Path(filename)
    suffix = path.suffix
    stem_limit = max(1, max_length - len(suffix))
    return f"{path.stem[:stem_limit]}{suffix}"


def _extract_html_text(path: Path) -> str:
    parser = _TextExtractor()
    parser.feed(path.read_text(encoding="utf-8"))
    return parser.text()


def _extract_pdf_text(path: Path) -> str:
    try:
        import pdfplumber
    except ImportError as exc:
        raise ValueError("缺少PDF解析依赖 pdfplumber") from exc

    parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            parts.append(page.extract_text() or "")
    return _compact_text("\n".join(parts))


def _extract_docx_text(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError("缺少DOCX解析依赖 python-docx") from exc

    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return _compact_text("\n".join(parts))


def _analyze_image(path: Path, filename: str, llm: Any | None) -> str:
    if llm is None:
        raise ValueError("图片附件需要可用的Vision模型")
    supports_multimodal = getattr(llm, "supports_multimodal", None)
    if callable(supports_multimodal) and not supports_multimodal():
        raise ValueError("当前LLM配置不支持图片分析，请切换Vision模型后重试")
    try:
        from crewai_files import FilePath
    except ImportError as exc:
        raise ValueError("缺少图片分析依赖 crewai-files") from exc

    prompt = (
        "请阅读这张需求附件图片，提取与产品需求相关的事实、界面元素、"
        "业务流程、数据字段、约束和待确认点。只输出中文摘要，不要编造图片中没有的信息。"
    )
    messages = [
        {
            "role": "user",
            "content": prompt,
            "files": {filename: FilePath(path=path)},
        }
    ]
    return str(llm.call(messages)).strip()


def _compact_text(value: str) -> str:
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in value.splitlines()]
    return "\n".join(line for line in lines if line)
